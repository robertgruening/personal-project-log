import datetime
import docx
from docx.shared import Pt
import json
from json import JSONEncoder

class Role():
    def __init__(self):
        self.Name = None
        self.Tasks = []

class Project():
    def __init__(self):
        self.Title = None
        self.StartDate = None
        self.EndDate = None
        self.Description = None
        self.CustomerName = None
        self.Roles = []
        self.Tags = []

class ProjectEncoder(JSONEncoder):
    def default(self, o):
        return o.__dict__
        
def convert_to_projects(json_projects):
    projects = []

    for json_project in json_projects:
        project = Project()
        project.Title = json_project.get('Title')
        project.StartDate = json_project.get('StartDate')
        project.EndDate = json_project.get('EndDate')
        project.Description = json_project.get('Description')
        project.CustomerName = json_project.get('CustomerName')

        for json_role in json_project.get('Roles'):
            role = Role()
            role.Name = json_role.get('Name')

            for json_task in json_role.get('Tasks'):
                role.Tasks.append(json_task)
            
            project.Roles.append(role)

        projects.append(project)

    return projects

def load_data():
    with open("projects.json", "r") as f:
        return convert_to_projects(json.load(f))

    return []

def export_format_a(doc, projects):
    table = doc.add_table(rows = len(projects), cols = 2)

    for i, project in enumerate(projects):
        row = table.rows[i].cells
        left_cell = row[0]
        left_cell.paragraphs.clear()
        left_cell_paragraph = left_cell.add_paragraph(style=None)

        right_cell = row[1]
        right_cell.paragraphs.clear()

        timespan = project.StartDate

        if project.EndDate is not None and\
            project.StartDate != project.EndDate:
            timespan += f" - {project.EndDate}"

        timespan_run = left_cell_paragraph.add_run(timespan)
        timespan_run.font.name = 'Arial'
        timespan_run.font.size = Pt(10)
        
        right_cell.add_paragraph(style=None)\
            .add_run(project.CustomerName)\
            .bold = True

        role_names = []

        for role in project.Roles:
            role_names.append(role.Name)

        run_role_names = right_cell.add_paragraph(style=None)\
            .add_run(', '.join(role_names))
        run_role_names.font.name = 'Arial'
        run_role_names.font.size = Pt(10)

        run_role_title = right_cell.add_paragraph(style=None)\
            .add_run(project.Title)
        run_role_title.font.name = 'Arial'
        run_role_title.font.size = Pt(10)

        run_description = right_cell.add_paragraph(style=None)\
            .add_run(project.Description)
        run_description.font.name = 'Arial'
        run_description.font.size = Pt(10)
        
        run_tags = right_cell.add_paragraph(style=None)\
            .add_run(', '.join(project.Tags))
        run_tags.font.name = 'Arial'
        run_tags.font.size = Pt(10)

    return doc

export_format_a(docx.Document(), load_data())\
    .save(f"{datetime.datetime.now().strftime("%Y-%m-%dT%H:%M:%S_Projekte.docx")}")