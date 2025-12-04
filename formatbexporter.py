import docx
from docx.shared import Pt

class FormatBExporter():
    def __init__(self):
        self.file_path = None

    def set_file_path(self, file_path:str):
        self.file_path = file_path

        return self

    def export(self, projects:list):
        doc = docx.Document()

        for i, project in enumerate(projects):
            role_names = []

            for role in project.Roles:
                role_names.append(role.Name)

            run_role_names = doc.add_paragraph(style=None)\
                .add_run(', '.join(role_names))
            run_role_names.font.name = 'Cambria'
            run_role_names.font.size = Pt(12)
            run_role_names.bold = True

            run_title = doc.add_paragraph(style=None)\
                .add_run(project.Title)
            run_title.font.name = 'Cambira'
            run_title.font.size = Pt(11)

            timespan = project.StartDate

            if project.EndDate is not None and\
                project.StartDate != project.EndDate:
                timespan += f" - {project.EndDate}"

            run_timespan = doc.add_paragraph(style=None)\
                .add_run(timespan)
            run_timespan.font.name = 'Cambira'
            run_timespan.font.size = Pt(11)

        doc.save(self.file_path)