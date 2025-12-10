import datetime
import docx

from project import Project

class TagsExporter():
    def __init__(self):
        self.file_path = None

    def set_file_path(self, file_path:str):
        self.file_path = file_path

        return self
    
    def _get_projects_by_tag(self, tag:str, projects:list):
        projects_with_tag = []

        for project in projects:
            if tag in project.Tags:
                projects_with_tag.append(project)

        return projects_with_tag
    
    def _get_timespan_keys(self, start_year:int, start_month:int, end_year:int, end_month:int):
        current_month = start_month
        current_year = start_year
        timespan_keys = []
        timespan_keys.append(f"{current_year}-{current_month:03d}")
            
        while not (current_month == end_month and \
            current_year == end_year):
            current_month += 1

            if current_month == 13:
                current_month = 1
                current_year += 1
            
            timespan_keys.append(f"{current_year}-{current_month:03d}")

        return timespan_keys
        
    def _get_project_timespan_keys(self, project:Project):
        start_month = int(project.StartDate.split('/')[0])
        start_year = int(project.StartDate.split('/')[1])
        end_month = None
        end_year = None

        if project.EndDate == 'heute':
            end_month = datetime.datetime.now().month
            end_year = datetime.datetime.now().year
        else:
            end_month = int(project.EndDate.split('/')[0])
            end_year = int(project.EndDate.split('/')[1])

        return self._get_timespan_keys(start_year, start_month, end_year, end_month)
    
    def _get_max_timespan(self, tag:str, projects:list):
        project_timespan_keys = []

        for project in self._get_projects_by_tag(tag, projects):
            for project_timespan_key in self._get_project_timespan_keys(project):
                project_timespan_keys.append(project_timespan_key)

        unique_project_timespan_keys = list(dict.fromkeys(project_timespan_keys))
        unique_project_timespan_keys.sort()

        min_year = int(unique_project_timespan_keys[0].split('-')[0])
        min_month = int(unique_project_timespan_keys[0].split('-')[1])
        max_year = int(unique_project_timespan_keys[len(unique_project_timespan_keys) - 1].split('-')[0])
        max_month = int(unique_project_timespan_keys[len(unique_project_timespan_keys) - 1].split('-')[1])

        return (f"{min_month}/{min_year:03d}", f"{max_month}/{max_year:03d}", len(self._get_timespan_keys(min_year, min_month, max_year, max_month)))
    
    def _get_min_timespan(self, tag:str, projects:list):
        project_timespan_keys = []

        for project in self._get_projects_by_tag(tag, projects):
            for project_timespan_key in self._get_project_timespan_keys(project):
                project_timespan_keys.append(project_timespan_key)

        unique_project_timespan_keys = list(dict.fromkeys(project_timespan_keys))

        return len(unique_project_timespan_keys)

    def _get_past_timespan(self, start_date:str):
        start_month = int(start_date.split('/')[0])
        start_year = int(start_date.split('/')[1])
        end_month = datetime.datetime.now().month
        end_year = datetime.datetime.now().year

        return len(self._get_timespan_keys(start_year, start_month, end_year, end_month)) - 1

    def export(self, projects:list):
        tags = []

        for i, project in enumerate(projects):
            for tag in project.Tags:
                if tag not in tags:
                    tags.append(tag)
        
        tags.sort()

        doc = docx.Document()

        run_tags = doc.add_paragraph(style=None)\
            .add_run('Tags')
        run_tags.bold = True

        table = doc.add_table(rows = len(tags) + 1, cols = 5)
        table.rows[0].cells[0].add_paragraph(style=None).add_run('Tag')
        table.rows[0].cells[1].add_paragraph(style=None).add_run('Zeitspanne')
        table.rows[0].cells[2].add_paragraph(style=None).add_run('Anz. Jahre (Zeitspanne)')
        table.rows[0].cells[3].add_paragraph(style=None).add_run('Anz. Jahre (Nutzung)')
        table.rows[0].cells[4].add_paragraph(style=None).add_run('Anz. Jahre (letzte Nutzung)')

        for i,tag in enumerate(tags):
            row = table.rows[i + 1].cells

            cell = row[0]
            cell.paragraphs.clear()
            cell_paragraph = cell.add_paragraph(style=None)
            cell_paragraph.add_run(tag)

            cell = row[1]
            cell.paragraphs.clear()
            cell_paragraph = cell.add_paragraph(style=None)
            max_timespan = self._get_max_timespan(tag, projects)
            cell_paragraph.add_run(f"{max_timespan[0]} - {max_timespan[1]}")

            cell = row[2]
            cell.paragraphs.clear()
            cell_paragraph = cell.add_paragraph(style=None)
            cell_paragraph.add_run(f"{(max_timespan[2]/12):.2f}")

            cell = row[3]
            cell.paragraphs.clear()
            cell_paragraph = cell.add_paragraph(style=None)
            min_timespan = self._get_min_timespan(tag, projects)
            cell_paragraph.add_run(f"{(min_timespan/12):.2f}")

            cell = row[4]
            cell.paragraphs.clear()
            cell_paragraph = cell.add_paragraph(style=None)
            past_timespan = self._get_past_timespan(max_timespan[1])
            cell_paragraph.add_run(f"{(past_timespan/12):.2f}")

        doc.save(self.file_path)