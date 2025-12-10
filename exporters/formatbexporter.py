import docx
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.text import WD_BREAK

class FormatBExporter():
    def __init__(self):
        self.file_path = None

    def set_file_path(self, file_path:str):
        self.file_path = file_path

        return self

    def export(self, projects:list):
        doc = docx.Document()

        for i, project in enumerate(projects):
            run_role_names = doc.add_paragraph(style=None)\
                .add_run(', '.join(project.RoleNames))
            run_role_names.font.name = 'Cambria'
            run_role_names.font.size = Pt(12)
            run_role_names.bold = True

            run_title = doc.add_paragraph(style=None)\
                .add_run(project.Title)
            run_title.font.name = 'Cambira'
            run_title.font.size = Pt(11)

            timespan = project.StartDate

            if project.EndDate is not None and\
                project.StartDate != project.EndDate:
                timespan += f" - {project.EndDate}"

            run_timespan = doc.add_paragraph(style=None)\
                .add_run(timespan)
            run_timespan.font.name = 'Cambira'
            run_timespan.font.size = Pt(11)

            doc.add_paragraph(style=None)

            table = doc.add_table(rows = 4, cols = 2)
            row = table.rows[0].cells
            left_cell = row[0]
            left_cell.paragraphs.clear()
            run_left_cell = left_cell.add_paragraph(style=None)\
                .add_run('KUNDE')
            run_left_cell.font.name = 'Cambira'
            run_left_cell.font.size = Pt(11)
            run_left_cell.bold = True

            right_cell = row[1]
            right_cell.paragraphs.clear()
            run_right_cell = right_cell.add_paragraph(style=None)\
                .add_run(f"{project.CustomerName}, {project.CustomerLocation}")
            run_right_cell.font.name = 'Cambira'
            run_right_cell.font.size = Pt(11)
            run_right_cell.bold = True

            row = table.rows[1].cells
            left_cell = row[0]
            left_cell.paragraphs.clear()
            run_left_cell = left_cell.add_paragraph(style=None)\
                .add_run('Branche')
            run_left_cell.font.name = 'Cambira'
            run_left_cell.font.size = Pt(11)
            run_left_cell.bold = True

            right_cell = row[1]
            right_cell.paragraphs.clear()
            run_right_cell = right_cell.add_paragraph(style=None)\
                .add_run(f"{project.IndustrySector}")
            run_right_cell.font.name = 'Cambira'
            run_right_cell.font.size = Pt(11)

            row = table.rows[2].cells
            left_cell = row[0]
            left_cell.paragraphs.clear()
            run_left_cell = left_cell.add_paragraph(style=None)\
                .add_run('Team- / Projektgröße')
            run_left_cell.font.name = 'Cambira'
            run_left_cell.font.size = Pt(11)
            run_left_cell.bold = True

            row = table.rows[3].cells
            left_cell = row[0]
            left_cell.paragraphs.clear()
            run_left_cell = left_cell.add_paragraph(style=None)\
                .add_run('Umfeld')
            run_left_cell.font.name = 'Cambira'
            run_left_cell.font.size = Pt(11)
            run_left_cell.bold = True
            run_left_cell = left_cell.add_paragraph(style=None)\
                .add_run('(Eingesetzte Methoden, \nProgrammiersprachen, \nWerkzeuge)')
            run_left_cell.font.name = 'Calibri'
            run_left_cell.font.size = Pt(8)

            right_cell = row[1]
            right_cell.paragraphs.clear()
            paragraph_left_cell = right_cell.add_paragraph(style=None)
            paragraph_left_cell.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            run_right_cell = paragraph_left_cell\
                .add_run(' | '.join(project.Tags))
            run_right_cell.font.name = 'Cambira'
            run_right_cell.font.size = Pt(11)

            doc.add_paragraph(style=None)

            paragraph_description = doc.add_paragraph(style=None)
            paragraph_description.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            run_description = paragraph_description\
                .add_run(project.Description)
            run_description.font.name = 'Calibri'
            run_description.font.size = Pt(11)

            doc.add_paragraph(style=None)

            paragraph_chapter_tasks = doc.add_paragraph(style=None)
            paragraph_chapter_tasks.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            run_chapter_tasks = paragraph_chapter_tasks\
                .add_run('AUFGABEN IM PROJEKT')
            run_chapter_tasks.font.name = 'Calibri'
            run_chapter_tasks.font.size = Pt(12)
            run_chapter_tasks.bold = True

            for task_group in project.TaskGroups:
                run_chapter_task_group = doc.add_paragraph(style=None)\
                    .add_run(task_group.Name)
                run_chapter_task_group.font.name = 'Arial'
                run_chapter_task_group.font.size = Pt(11)
                run_chapter_task_group.bold = True

                for task in task_group.Tasks:
                    run_chapter_task = doc.add_paragraph(style='List Bullet')\
                        .add_run(task)
                    run_chapter_task.font.name = 'Calibri'
                    run_chapter_task.font.size = Pt(11)


            paragraph_new_page = doc.add_paragraph(style=None)
            run_new_page = paragraph_new_page.add_run()
            run_new_page.add_break(WD_BREAK.PAGE)

        doc.save(self.file_path)