import docx
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.text import WD_BREAK

class FormatCExporter():
    def __init__(self):
        self.file_path = None

    def set_file_path(self, file_path:str):
        self.file_path = file_path

        return self

    def export(self, projects:list):
        doc = docx.Document()

        for i, project in enumerate(projects):
            chapter_project = doc.add_paragraph(style='List Number')

            run_title = chapter_project.add_run(project.Title)
            run_title.font.name = 'Calibri'
            run_title.font.size = Pt(11)
            run_title.bold = True
            chapter_project.add_run().add_break(WD_BREAK.LINE)

            run_industry_sector = chapter_project.add_run(f"Branche:\t\t\t\t{project.IndustrySector}")
            run_industry_sector.font.name = 'Calibri'
            run_industry_sector.font.size = Pt(11)
            chapter_project.add_run().add_break(WD_BREAK.LINE)

            timespan = project.StartDate

            if project.EndDate is not None and\
                project.StartDate != project.EndDate:
                timespan += f" - {project.EndDate}"

            run_industry_sector = chapter_project.add_run(f"Zeitraum:\t\t\t\t{timespan}")
            run_industry_sector.font.name = 'Calibri'
            run_industry_sector.font.size = Pt(11)
            chapter_project.add_run().add_break(WD_BREAK.LINE)

            run_industry_sector = chapter_project.add_run(f"Rolle:\t\t\t\t{', '.join(project.RoleNames)}")
            run_industry_sector.font.name = 'Calibri'
            run_industry_sector.font.size = Pt(11)
            chapter_project.add_run().add_break(WD_BREAK.LINE)

            run_industry_sector = chapter_project.add_run(f"Projektbeschreibung:\t{project.Description}")
            run_industry_sector.font.name = 'Calibri'
            run_industry_sector.font.size = Pt(11)
            chapter_project.add_run().add_break(WD_BREAK.LINE)

            tasks = []

            for task_group in project.TaskGroups:
                for task in task_group.Tasks:
                    tasks.append(task)

            for i,task in enumerate(tasks):
                if i == 0:
                    run_industry_sector = chapter_project.add_run(f"Aufgaben:\t\t\t\t{task}")
                    run_industry_sector.font.name = 'Calibri'
                    run_industry_sector.font.size = Pt(11)
                else:
                    run_industry_sector = chapter_project.add_run(f"\t\t\t\t\t{task}")
                    run_industry_sector.font.name = 'Calibri'
                    run_industry_sector.font.size = Pt(11)
                if i < len(tasks):
                    run_industry_sector.add_break(WD_BREAK.LINE)

            chapter_project.add_run().add_break(WD_BREAK.LINE)

            run_industry_sector = chapter_project.add_run(f"Sprachen/Technologien:\t\t\t\t{', '.join(project.Tags)}")
            run_industry_sector.font.name = 'Calibri'
            run_industry_sector.font.size = Pt(11)
            chapter_project.add_run().add_break(WD_BREAK.LINE)

            run_industry_sector = chapter_project.add_run(f"Methodik/Vorgehen:\t\t\t\t{', '.join(project.Tags)}")
            run_industry_sector.font.name = 'Calibri'
            run_industry_sector.font.size = Pt(11)

            doc.add_paragraph(style=None)
            doc.add_paragraph(style=None)

        doc.save(self.file_path)