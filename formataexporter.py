import docx
from docx.shared import Pt

class FormatAExporter():
    def __init__(self):
        self.file_path = None

    def set_file_path(self, file_path:str):
        self.file_path = file_path

        return self

    def export(self, projects:list):
        doc = docx.Document()
        table = doc.add_table(rows = len(projects), cols = 2)

        for i, project in enumerate(projects):
            row = table.rows[i].cells
            left_cell = row[0]
            left_cell.paragraphs.clear()
            left_cell_paragraph = left_cell.add_paragraph(style=None)

            right_cell = row[1]
            right_cell.paragraphs.clear()

            timespan = project.StartDate

            if project.EndDate is not None and\
                project.StartDate != project.EndDate:
                timespan += f" - {project.EndDate}"

            timespan_run = left_cell_paragraph.add_run(timespan)
            timespan_run.font.name = 'Arial'
            timespan_run.font.size = Pt(10)
            
            right_cell.add_paragraph(style=None)\
                .add_run(project.CustomerName)\
                .bold = True

            run_role_names = right_cell.add_paragraph(style=None)\
                .add_run(', '.join(project.RoleNames))
            run_role_names.font.name = 'Arial'
            run_role_names.font.size = Pt(10)

            run_title = right_cell.add_paragraph(style=None)\
                .add_run(project.Title)
            run_title.font.name = 'Arial'
            run_title.font.size = Pt(10)

            run_description = right_cell.add_paragraph(style=None)\
                .add_run(project.Description)
            run_description.font.name = 'Arial'
            run_description.font.size = Pt(10)
            
            run_tags = right_cell.add_paragraph(style=None)\
                .add_run(', '.join(project.Tags))
            run_tags.font.name = 'Arial'
            run_tags.font.size = Pt(10)

        doc.save(self.file_path)