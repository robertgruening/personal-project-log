import datetime
import docx

class TagsExporter():
    def __init__(self):
        self.file_path = None

    def set_file_path(self, file_path:str):
        self.file_path = file_path

        return self
    
    def _get_min_timespan(self, tag:str, projects:list):
        projects_with_tag = []

        for project in projects:
            if tag in project.Tags:
                projects_with_tag.append(project)

        for project_with_tag in projects_with_tag:
            pass

        timespan = 0

        return 0
    
    def _get_max_timespan(self, tag:str, projects:list):
        projects_with_tag = []

        for project in projects:
            if tag in project.Tags:
                projects_with_tag.append(project)

        min_month = int(projects_with_tag[0].StartDate.split('/')[0])
        min_year = int(projects_with_tag[0].StartDate.split('/')[1])

        max_month = None
        max_year = None

        if projects_with_tag[0].EndDate == 'heute':
            max_month = datetime.datetime.now().month
            max_year = datetime.datetime.now().year
        else:
            max_month = int(projects_with_tag[0].EndDate.split('/')[0])
            max_year = int(projects_with_tag[0].EndDate.split('/')[1])

        for project_with_tag in projects_with_tag:
            start_month = int(project_with_tag.StartDate.split('/')[0])
            start_year = int(project_with_tag.StartDate.split('/')[1])

            if start_year < min_year:
                min_year = start_year
                min_month = start_month
            elif start_year == min_year and \
                start_month < min_month:
                min_year = start_year
                min_month = start_month
            
            if project_with_tag.EndDate == 'heute':
                max_year = datetime.datetime.now().year
                max_month = datetime.datetime.now().month
                continue

            end_month = int(project_with_tag.EndDate.split('/')[0])
            end_year = int(project_with_tag.EndDate.split('/')[1])

            if end_year > max_year:
                max_year = end_year
                max_month = end_month
            elif end_year == max_year and \
                end_month > max_month:
                max_year = end_year
                max_month = end_month

        return (f"{min_month}/{min_year}", f"{max_month}/{max_year}")

    def export(self, projects:list):
        tags = []

        for i, project in enumerate(projects):
            for tag in project.Tags:
                if tag not in tags:
                    tags.append(tag)
        
        tags.sort()

        doc = docx.Document()

        run_tags = doc.add_paragraph(style=None)\
            .add_run('Tags')
        run_tags.bold = True

        table = doc.add_table(rows = len(tags) + 1, cols = 3)
        table.rows[0].cells[0].add_paragraph(style=None).add_run('Tag')
        table.rows[0].cells[1].add_paragraph(style=None).add_run('max. Zeitspanne')
        table.rows[0].cells[2].add_paragraph(style=None).add_run('min. Zeitspanne')

        for i,tag in enumerate(tags):
            row = table.rows[i + 1].cells

            cell = row[0]
            cell.paragraphs.clear()
            cell_paragraph = cell.add_paragraph(style=None)
            cell_paragraph.add_run(tag)

            cell = row[1]
            cell.paragraphs.clear()
            cell_paragraph = cell.add_paragraph(style=None)
            max_timespan = self._get_max_timespan(tag, projects)
            cell_paragraph.add_run(f"{max_timespan[0]} - {max_timespan[1]}")

            cell = row[2]
            cell.paragraphs.clear()
            cell_paragraph = cell.add_paragraph(style=None)
            cell_paragraph.add_run(str(self._get_min_timespan(tag, projects)))

        doc.save(self.file_path)